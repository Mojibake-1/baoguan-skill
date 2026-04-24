#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Fill Excel or Word customs declaration templates from JSON data."""

from __future__ import annotations

import argparse
import json
import re
import sys
from copy import copy, deepcopy
from pathlib import Path
from typing import Any


PLACEHOLDER_RE = re.compile(r"{{\s*([^{}]+?)\s*}}")
FULL_PLACEHOLDER_RE = re.compile(r"^\s*{{\s*([^{}]+?)\s*}}\s*$")
ARRAY_TOKEN_RE = re.compile(r"^([A-Za-z_][\w-]*)\[\]\.(.+)$")
INDEXED_PART_RE = re.compile(r"^([A-Za-z_][\w-]*)(?:\[(\d+)\])?$")
MISSING = object()


class Renderer:
    def __init__(self, data: dict[str, Any], missing: str, strict: bool) -> None:
        self.data = data
        self.missing = missing
        self.strict = strict
        self.missing_tokens: set[str] = set()

    def render(self, text: str, array_context: dict[str, Any] | None = None, preserve_type: bool = False) -> Any:
        if "{{" not in text:
            return text

        full_match = FULL_PLACEHOLDER_RE.match(text)
        if full_match:
            token = full_match.group(1).strip()
            value = self.resolve(token, array_context or {})
            if value is MISSING:
                return self.missing_value(token)
            if preserve_type and is_scalar(value):
                return value
            return stringify(value)

        def replace(match: re.Match[str]) -> str:
            token = match.group(1).strip()
            value = self.resolve(token, array_context or {})
            if value is MISSING:
                return str(self.missing_value(token))
            return stringify(value)

        return PLACEHOLDER_RE.sub(replace, text)

    def resolve(self, token: str, array_context: dict[str, Any]) -> Any:
        array_match = ARRAY_TOKEN_RE.match(token)
        if array_match:
            array_name, item_path = array_match.groups()
            if array_name not in array_context:
                return MISSING
            return get_path(array_context[array_name], item_path)

        value = get_path(self.data, token)
        return value

    def missing_value(self, token: str) -> str:
        self.missing_tokens.add(token)
        if self.missing == "keep":
            return "{{" + token + "}}"
        if self.missing == "blank":
            return ""
        return "\u5f85\u8865\u5145:" + token

    def raise_if_missing(self) -> None:
        if self.strict and self.missing_tokens:
            tokens = ", ".join(sorted(self.missing_tokens))
            raise ValueError(f"Unresolved placeholders: {tokens}")


def is_scalar(value: Any) -> bool:
    return value is None or isinstance(value, (str, int, float, bool))


def stringify(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def get_path(root: Any, path: str) -> Any:
    if isinstance(root, dict) and path in root:
        return root[path]

    current = root
    for part in path.split("."):
        part = part.strip()
        if not part:
            continue

        indexed = INDEXED_PART_RE.match(part)
        if indexed:
            key, index = indexed.groups()
        else:
            key, index = part, None

        if isinstance(current, dict):
            if key not in current:
                return MISSING
            current = current[key]
        else:
            return MISSING

        if index is not None:
            if not isinstance(current, list):
                return MISSING
            position = int(index)
            if position >= len(current):
                return MISSING
            current = current[position]

    return current


def load_json(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8-sig") as handle:
        data = json.load(handle)
    if not isinstance(data, dict):
        raise ValueError("Data JSON must contain an object at the top level.")
    return data


def find_array_names(text: str) -> set[str]:
    names: set[str] = set()
    for token in PLACEHOLDER_RE.findall(text):
        match = ARRAY_TOKEN_RE.match(token.strip())
        if match:
            names.add(match.group(1))
    return names


def fill_xlsx(template: Path, output: Path, renderer: Renderer) -> None:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("openpyxl is required to fill .xlsx templates.") from exc

    workbook = load_workbook(template)
    for sheet in workbook.worksheets:
        expand_xlsx_rows(sheet, renderer)
        for row in sheet.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    cell.value = renderer.render(cell.value, preserve_type=True)

    workbook.save(output)


def expand_xlsx_rows(sheet: Any, renderer: Renderer) -> None:
    repeat_rows: list[tuple[int, str]] = []

    for row_index in range(1, sheet.max_row + 1):
        names: set[str] = set()
        for column_index in range(1, sheet.max_column + 1):
            value = sheet.cell(row_index, column_index).value
            if isinstance(value, str):
                names.update(find_array_names(value))
        if names:
            if len(names) > 1:
                raise ValueError(
                    f"Sheet {sheet.title!r} row {row_index} mixes repeat arrays: {', '.join(sorted(names))}"
                )
            repeat_rows.append((row_index, next(iter(names))))

    for row_index, array_name in reversed(repeat_rows):
        values = get_path(renderer.data, array_name)
        items = values if isinstance(values, list) else []
        count = max(1, len(items))

        if count > 1:
            sheet.insert_rows(row_index + 1, amount=count - 1)
            for offset in range(1, count):
                copy_xlsx_row(sheet, row_index, row_index + offset)

        for offset in range(count):
            item = items[offset] if offset < len(items) else {}
            render_xlsx_row(sheet, row_index + offset, renderer, {array_name: item})


def copy_xlsx_row(sheet: Any, source_row: int, target_row: int) -> None:
    source_dimensions = sheet.row_dimensions[source_row]
    target_dimensions = sheet.row_dimensions[target_row]
    target_dimensions.height = source_dimensions.height
    target_dimensions.hidden = source_dimensions.hidden

    for column_index in range(1, sheet.max_column + 1):
        source = sheet.cell(source_row, column_index)
        target = sheet.cell(target_row, column_index)
        target.value = source.value
        if source.has_style:
            target._style = copy(source._style)
        if source.number_format:
            target.number_format = source.number_format
        if source.alignment:
            target.alignment = copy(source.alignment)
        if source.font:
            target.font = copy(source.font)
        if source.fill:
            target.fill = copy(source.fill)
        if source.border:
            target.border = copy(source.border)
        if source.protection:
            target.protection = copy(source.protection)
        if source.comment:
            target.comment = copy(source.comment)


def render_xlsx_row(sheet: Any, row_index: int, renderer: Renderer, array_context: dict[str, Any]) -> None:
    for column_index in range(1, sheet.max_column + 1):
        cell = sheet.cell(row_index, column_index)
        if isinstance(cell.value, str):
            cell.value = renderer.render(cell.value, array_context=array_context, preserve_type=True)


def fill_docx(template: Path, output: Path, renderer: Renderer) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required to fill .docx templates.") from exc

    document = Document(template)
    for part in iter_docx_parts(document):
        for table in part.tables:
            expand_docx_table(table, renderer)
        for paragraph in part.paragraphs:
            render_docx_paragraph(paragraph, renderer, {})
        for table in part.tables:
            render_docx_table(table, renderer, {})

    document.save(output)


def iter_docx_parts(document: Any) -> list[Any]:
    parts = [document]
    for section in document.sections:
        parts.extend(
            [
                section.header,
                section.footer,
                section.first_page_header,
                section.first_page_footer,
                section.even_page_header,
                section.even_page_footer,
            ]
        )
    return parts


def expand_docx_table(table: Any, renderer: Renderer) -> None:
    try:
        from docx.table import _Row
    except ImportError as exc:
        raise RuntimeError("python-docx internals changed; cannot duplicate table rows.") from exc

    for row in reversed(list(table.rows)):
        row_text = "\n".join(cell.text for cell in row.cells)
        names = find_array_names(row_text)
        if not names:
            continue
        if len(names) > 1:
            raise ValueError(f"A Word table row mixes repeat arrays: {', '.join(sorted(names))}")

        array_name = next(iter(names))
        values = get_path(renderer.data, array_name)
        items = values if isinstance(values, list) else []
        count = max(1, len(items))
        row_elements = [row._tr]
        insert_after = row._tr

        for _ in range(1, count):
            new_row = deepcopy(row._tr)
            insert_after.addnext(new_row)
            insert_after = new_row
            row_elements.append(new_row)

        for offset, row_element in enumerate(row_elements):
            item = items[offset] if offset < len(items) else {}
            render_docx_row(_Row(row_element, table), renderer, {array_name: item})


def render_docx_table(table: Any, renderer: Renderer, array_context: dict[str, Any]) -> None:
    for row in table.rows:
        render_docx_row(row, renderer, array_context)
        for cell in row.cells:
            for nested_table in cell.tables:
                render_docx_table(nested_table, renderer, array_context)


def render_docx_row(row: Any, renderer: Renderer, array_context: dict[str, Any]) -> None:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            render_docx_paragraph(paragraph, renderer, array_context)
        for table in cell.tables:
            render_docx_table(table, renderer, array_context)


def render_docx_paragraph(paragraph: Any, renderer: Renderer, array_context: dict[str, Any]) -> None:
    if "{{" not in paragraph.text:
        return

    for run in paragraph.runs:
        if "{{" in run.text:
            run.text = renderer.render(run.text, array_context=array_context)

    if "{{" in paragraph.text:
        replacement = renderer.render(paragraph.text, array_context=array_context)
        if paragraph.runs:
            paragraph.runs[0].text = replacement
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(replacement)


def parse_args(argv: list[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Fill an Excel or Word customs declaration template from JSON data.")
    parser.add_argument("--template", required=True, type=Path, help="Path to .xlsx or .docx template.")
    parser.add_argument("--data", required=True, type=Path, help="Path to JSON data.")
    parser.add_argument("--output", required=True, type=Path, help="Path to write the filled document.")
    parser.add_argument(
        "--missing",
        choices=["marker", "keep", "blank"],
        default="marker",
        help="How to render unresolved placeholders. Default: marker.",
    )
    parser.add_argument("--strict", action="store_true", help="Fail when any placeholder cannot be resolved.")
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv or sys.argv[1:])
    template = args.template
    output = args.output

    if not template.exists():
        raise FileNotFoundError(f"Template not found: {template}")
    if not args.data.exists():
        raise FileNotFoundError(f"Data file not found: {args.data}")

    output.parent.mkdir(parents=True, exist_ok=True)
    renderer = Renderer(load_json(args.data), args.missing, args.strict)

    suffix = template.suffix.lower()
    if suffix == ".xlsx":
        fill_xlsx(template, output, renderer)
    elif suffix == ".docx":
        fill_docx(template, output, renderer)
    else:
        raise ValueError("Template must be .xlsx or .docx")

    renderer.raise_if_missing()
    if renderer.missing_tokens and not args.strict:
        print("Unresolved placeholders:", ", ".join(sorted(renderer.missing_tokens)), file=sys.stderr)
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:
        print(f"error: {exc}", file=sys.stderr)
        raise SystemExit(1)
